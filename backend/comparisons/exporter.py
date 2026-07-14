from __future__ import annotations

import html
import json
from pathlib import Path

from backend.comparisons.service import to_markdown


MEDIA = {
    "markdown": "text/markdown; charset=utf-8",
    "json": "application/json",
    "html": "text/html; charset=utf-8",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


class ComparisonExporter:
    def save_all(self, report: dict, directory: Path, comparison_id: str) -> dict[str, str]:
        directory.mkdir(parents=True, exist_ok=True)
        paths = {fmt: directory / f"{comparison_id}_comparison.{('md' if fmt == 'markdown' else fmt)}" for fmt in MEDIA}
        markdown = to_markdown(report)
        paths["markdown"].write_text(markdown, encoding="utf-8")
        paths["json"].write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        html_text = self._html(report)
        paths["html"].write_text(html_text, encoding="utf-8")
        self._pdf(html_text, paths["pdf"])
        self._docx(report, paths["docx"])
        return {key: str(value) for key, value in paths.items()}

    def _html(self, report: dict) -> str:
        matrix = ["<table><thead><tr><th>Dimension</th>" + "".join(f"<th>{html.escape(p['title'])}</th>" for p in report["source_papers"]) + "</tr></thead><tbody>"]
        for row in report["matrix"]:
            matrix.append(f"<tr><th>{html.escape(row['dimension'])}</th>" + "".join(f"<td>{html.escape(cell['summary'])}<small>{html.escape(', '.join(cell['evidence_ids']))}</small></td>" for cell in row["cells"]) + "</tr>")
        matrix.append("</tbody></table>")
        synthesis = "".join(f"<h2>{html.escape(name.replace('_',' ').title())}</h2><p>{html.escape(value['content'])}</p><small>{html.escape(', '.join(value['evidence_ids']))}</small>" for name, value in report["synthesis"].items())
        return "<!doctype html><html><head><meta charset='utf-8'><style>body{font:15px/1.55 system-ui;max-width:1200px;margin:2rem auto;padding:0 2rem;color:#17251e}table{border-collapse:collapse;width:100%}th,td{border:1px solid #dce5df;padding:.7rem;vertical-align:top}th{background:#f1f5f2}small{display:block;margin-top:.5rem;color:#557064}</style></head><body>" + f"<h1>{html.escape(report['title'])}</h1><p><strong>Focus:</strong> {html.escape(report['focus'])}</p>" + "".join(matrix) + synthesis + "</body></html>"

    def _pdf(self, html_text: str, target: Path) -> None:
        import pymupdf
        story = pymupdf.Story(html=html_text)
        writer = pymupdf.DocumentWriter(str(target))
        mediabox = pymupdf.paper_rect("a4-l")
        where = mediabox + (30, 35, -30, -35)
        more = True
        while more:
            device = writer.begin_page(mediabox)
            more, _ = story.place(where)
            story.draw(device)
            writer.end_page()
        writer.close()

    def _docx(self, report: dict, target: Path) -> None:
        from docx import Document
        document = Document()
        document.add_heading(report["title"], 0)
        document.add_paragraph(f"Focus: {report['focus']}")
        table = document.add_table(rows=1, cols=1 + len(report["source_papers"]))
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Dimension"
        for index, paper in enumerate(report["source_papers"], 1):
            table.rows[0].cells[index].text = paper["title"]
        for row in report["matrix"]:
            cells = table.add_row().cells
            cells[0].text = row["dimension"]
            for index, cell in enumerate(row["cells"], 1):
                cells[index].text = cell["summary"] + "\n" + ", ".join(cell["evidence_ids"])
        document.add_heading("Synthesis", level=1)
        for name, section in report["synthesis"].items():
            document.add_heading(name.replace("_", " ").title(), level=2)
            document.add_paragraph(section["content"])
            document.add_paragraph("Evidence: " + ", ".join(section["evidence_ids"]))
        document.core_properties.comments = json.dumps({"schema_version": "paper-comparison-v1"})
        document.save(target)
