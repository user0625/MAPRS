from __future__ import annotations

import html
import json
from pathlib import Path
from typing import Literal

from backend.schemas.report import FinalReport


class ArtifactExporter:
  MEDIA = {
    "markdown": "text/markdown; charset=utf-8",
    "json": "application/json",
    "html": "text/html; charset=utf-8",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
  }

  def get_or_create(self, record, format: Literal["markdown", "json", "html", "pdf", "docx"]):
    markdown = Path(record.report_path)
    base = markdown.with_suffix("")
    json_path = markdown.with_suffix(".json")
    if format == "markdown":
      return markdown, self.MEDIA[format]
    if not json_path.exists():
      raise ValueError("Structured report JSON is missing.")
    if format == "json":
      return json_path, self.MEDIA[format]
    target = Path(f"{base}.{format}")
    if target.exists() and target.stat().st_mtime >= json_path.stat().st_mtime:
      return target, self.MEDIA[format]
    report = FinalReport.model_validate_json(json_path.read_text(encoding="utf-8"))
    if format == "html":
      target.write_text(self._html(report), encoding="utf-8")
    elif format == "pdf":
      self._pdf(report, target)
    else:
      self._docx(report, target)
    return target, self.MEDIA[format]

  def _html(self, report: FinalReport) -> str:
    parts = ["<!doctype html><html><head><meta charset='utf-8'>",
      "<style>body{font:16px/1.65 system-ui,sans-serif;max-width:900px;margin:3rem auto;padding:0 2rem;color:#17202a}h1,h2{line-height:1.25}.warning{background:#fff4ce;padding:1rem;border-left:4px solid #c58b00}.evidence{color:#53606d;font-size:.9em}</style></head><body>",
      f"<h1>{html.escape(report.title)}</h1>"]
    if report.paper_title:
      parts.append(f"<p><strong>Paper:</strong> {html.escape(report.paper_title)}</p>")
    if report.warning:
      parts.append(f"<p class='warning'>{html.escape(report.warning)}</p>")
    for section in report.sorted_sections():
      content = "<br>".join(html.escape(section.content).splitlines())
      parts.extend([f"<h2>{html.escape(section.title)}</h2>", f"<p>{content}</p>"])
      if section.evidence_ids:
        parts.append(f"<p class='evidence'>Evidence: {html.escape(', '.join(section.evidence_ids))}</p>")
    parts.append("</body></html>")
    return "".join(parts)

  def _pdf(self, report: FinalReport, target: Path) -> None:
    import pymupdf
    story = pymupdf.Story(html=self._html(report))
    writer = pymupdf.DocumentWriter(str(target))
    mediabox = pymupdf.paper_rect("a4")
    where = mediabox + (36, 42, -36, -42)
    more = True
    while more:
      device = writer.begin_page(mediabox)
      more, _ = story.place(where)
      story.draw(device)
      writer.end_page()
    writer.close()

  def _docx(self, report: FinalReport, target: Path) -> None:
    from docx import Document
    document = Document()
    document.add_heading(report.title, 0)
    if report.paper_title:
      document.add_paragraph(f"Paper: {report.paper_title}")
    if report.warning:
      document.add_paragraph(f"⚠ {report.warning}")
    for section in report.sorted_sections():
      document.add_heading(section.title, level=1)
      document.add_paragraph(section.content)
      if section.evidence_ids:
        document.add_paragraph("Evidence: " + ", ".join(section.evidence_ids))
    document.core_properties.comments = json.dumps({"format": "phase-d"})
    document.save(target)
